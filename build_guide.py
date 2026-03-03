from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ─── Colour palette ──────────────────────────────────────────────────────────
BLACK   = RGBColor(0x0c, 0x0d, 0x0f)
ORANGE  = RGBColor(0xf5, 0xa6, 0x23)
WHITE   = RGBColor(0xff, 0xff, 0xff)
DARK    = RGBColor(0x17, 0x19, 0x1d)
LIGHT   = RGBColor(0xf0, 0xed, 0xe6)
GREY    = RGBColor(0x44, 0x44, 0x44)
LGREY   = RGBColor(0xf5, 0xf5, 0xf5)
GREEN   = RGBColor(0x22, 0xc5, 0x5e)
RED     = RGBColor(0xe5, 0x4d, 0x4d)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def rgb_hex(rgb: RGBColor) -> str:
    # RGBColor is a subclass of tuple: (r, g, b)
    return str(rgb).upper()

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex(rgb))
    tcPr.append(shd)

def set_para_bg(para, rgb: RGBColor):
    """Set paragraph shading (for code blocks)."""
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex(rgb))
    pPr.append(shd)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Pt(0)
    set_para_bg(p, DARK)
    run = p.add_run(f'  {text}')
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = ORANGE

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREY

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Pt(16)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Pt(18 + level * 14)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Pt(10)
        set_para_bg(p, RGBColor(0xf0, 0xf0, 0xf0))
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1e, 0x1e, 0x6e)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'F5A623')
    pBdr.append(bottom)
    pPr.append(pBdr)

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, DARK)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(9.5)
        run.font.color.rgb = ORANGE
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            if ri % 2 == 0:
                set_cell_bg(cell, LGREY)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size      = Pt(9.5)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    # column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# Orange background block via table trick
cov = doc.add_table(rows=1, cols=1)
cov.alignment = WD_TABLE_ALIGNMENT.CENTER
c = cov.rows[0].cells[0]
set_cell_bg(c, DARK)
c.width = Inches(6.27)

for txt, sz, bold, color in [
    ('🏀  CourtIQ', 36, True, ORANGE),
    ('חוברת מקצועית מלאה', 18, True, LIGHT),
    ('', 6, False, LIGHT),
    ('סיכום הפרויקט  ·  הסברים מעמיקים  ·  מפת דרכים', 11, False, RGBColor(0xa0,0xa0,0xa0)),
    ('', 6, False, LIGHT),
    ('מרץ 2026', 10, False, RGBColor(0x80,0x80,0x80)),
]:
    p   = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(txt)
    run.bold           = bold
    run.font.size      = Pt(sz)
    run.font.color.rgb = color

doc.add_paragraph()

# ─── TOC ─────────────────────────────────────────────────────────────────────
heading1('תוכן עניינים')

toc_items = [
    ('1', 'מה בנינו — סקירה כללית'),
    ('2', 'הסטאק הטכנולוגי — כל טכנולוגיה מוסברת'),
    ('3', 'מבנה הפרויקט — קבצים ותפקידיהם'),
    ('4', 'מסד הנתונים — הטבלאות והמשמעות שלהן'),
    ('5', 'האינטגרציה עם Claude AI'),
    ('6', 'כל תכונה שנבנתה — הסבר מפורט'),
    ('7', 'מה חסר — ניתוח מלא'),
    ('8', 'מפת דרכים לאפליקציה המושלמת'),
    ('9', 'מילון מונחים'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'  {num}.  {title}')
    run.font.size      = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x5e)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — What We Built
# ══════════════════════════════════════════════════════════════════════════════
heading1('1.  מה בנינו — סקירה כללית')
divider()

body('CourtIQ היא פלטפורמת אימון כדורסל מבוססת בינה מלאכותית.')
body('')
heading2('המטרה')
body('לתת לשחקני כדורסל בכל הרמות תוכנית אימון אישית, מעקב אחר התקדמות וניתוח ביצועים — הכל דרך AI.')

heading2('מה נבנה בפועל')
make_table(
    ['מה', 'קובץ / שירות', 'תיאור'],
    [
        ['דף נחיתה שיווקי', 'index.html', 'מוכר את המוצר — תמחור, תכונות, CTA'],
        ['דשבורד ראשי',     'app.html',   '10 פאנלים — האפליקציה עצמה'],
        ['Backend',         'Supabase',   'מסד נתונים, אימות, Edge Function'],
        ['AI Coach',        'Claude API', 'ניתוח שבועי + צ\'אט מאמן'],
    ],
    col_widths=[1.4, 1.5, 3.3]
)

heading2('שלבי הפיתוח — לפי Git History')
make_table(
    ['Commit', 'מה נעשה'],
    [
        ['Initial release',       'גרסה ראשונה — CourtIQ בסיסי'],
        ['UX redesign',           'פיצול index/app, עיצוב מלא, מערכת תמחור'],
        ['Supabase integration',  'חיבור DB, Auth, Edge Function'],
        ['SQL helpers',           'כלי עזר להקמת מסד הנתונים (sql.html, trigger_only.sql)'],
    ],
    col_widths=[2.2, 4.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — Tech Stack
# ══════════════════════════════════════════════════════════════════════════════
heading1('2.  הסטאק הטכנולוגי')
divider()

# ─── Frontend ────────────────────────────────────────────────────────────────
heading2('Frontend — צד הלקוח')

heading3('HTML5')
body('מה זה: שפת המבנה של כל דף ווב. כמו שלד של בניין.')
body('איך השתמשנו: כל האפליקציה נמצאת בשני קבצי HTML גדולים — index.html ו-app.html.')
body('למה ללא Framework: בחרנו Vanilla JS (ללא React/Vue) כי רצינו להתחיל מהיר ללא כלי build מורכבים.')

heading3('CSS3')
body('מה זה: שפת העיצוב — צבעים, פונטים, פריסה, אנימציות.')
body('איך השתמשנו:')
bullet('CSS Variables (משתנים גלובליים) לניהול הצבעים')
code_block([
    '--c-bg:      #0c0d0f   /* רקע כהה */',
    '--c-accent:  #f5a623   /* כתום-ענבר */',
    '--c-text:    #f0ede6   /* טקסט בהיר */',
    '--c-surface: #17191d   /* משטח כרטיסים */',
])
bullet('Flexbox ו-Grid לפריסה רספונסיבית')
bullet('Animations ו-Transitions לחווית משתמש חלקה')
bullet('Breakpoints: Desktop 1024px+ / Tablet 768px / Mobile <768px')

heading3('JavaScript (Vanilla)')
body('מה זה: שפת התכנות שמביאה את הדף לחיים — לוגיקה, אירועים, קריאות API.')
body('איך השתמשנו:')
bullet('State management דרך Object גלובלי')
code_block(['const APP = { user, profile, currentWeek, sessions }'])
bullet('Event listeners על כפתורים ותפריטים')
bullet('Async/Await לקריאות רשת')
bullet('DOM Manipulation לשינוי התצוגה')

heading3('Chart.js')
body('מה זה: ספריית גרפים מוכנה לשימוש.')
body('איך השתמשנו — 4 גרפים בפאנל Analytics:')
make_table(
    ['גרף', 'סוג', 'מה מציג'],
    [
        ['1', 'קו (Line)',     'אחוז זריקה לאורך הזמן'],
        ['2', 'עמודות (Bar)', 'השוואה בין שבועות'],
        ['3', 'קו',           'קפיצה אנכית (vertical)'],
        ['4', 'קו',           'זמן ספרינט'],
    ],
    col_widths=[0.5, 1.5, 4.2]
)

# ─── Backend ─────────────────────────────────────────────────────────────────
heading2('Backend — צד השרת')

heading3('Supabase')
body('מה זה: Backend-as-a-Service — שירות שנותן לך מסד נתונים, אימות, ו-API תוך דקות, ללא צורך לכתוב שרת.')
body('למה Supabase:')
bullet('PostgreSQL אמיתי (לא NoSQL)')
bullet('Row Level Security מובנה')
bullet('Edge Functions מבוסס Deno')
bullet('חינמי לפרויקטים קטנים')
bullet('דאשבורד ויזואלי לניהול')

make_table(
    ['שירות', 'מה זה', 'מה עשינו'],
    [
        ['Supabase Auth',           'ניהול משתמשים',          'הרשמה, כניסה, session'],
        ['Supabase Database',       'PostgreSQL בענן',         'שמירת פרופילים, אימונים, שבועות'],
        ['Supabase Edge Functions', 'שרת קוד בענן (Deno)',    'Proxy מאובטח לClaude API'],
    ],
    col_widths=[2.0, 2.0, 2.2]
)

heading3('Supabase Auth — איך זה עובד')
code_block([
    'משתמש → מזין email/password → Supabase Auth → מחזיר JWT Token',
    'JWT Token → נשמר בBrowser → נשלח בכל בקשה → Supabase מאמת',
])
body('JWT Token: מחרוזת מוצפנת שמזהה את המשתמש — כמו תעודת זהות דיגיטלית.')

heading3('Row Level Security (RLS)')
body('מה זה: מדיניות אבטחה ישירות במסד הנתונים.')
body('הבעיה שהיא פותרת: בלי RLS, אם יש bug בקוד, משתמש יכול לראות נתוני משתמשים אחרים.')
body('איך עובד:')
code_block([
    'create policy "sessions: select own"',
    '  on training_sessions',
    '  for select',
    '  using (auth.uid() = user_id);',
    '',
    '-- = "משתמש יכול לראות רק שורות שה-user_id שלהן = ה-ID שלו"',
])

# ─── AI ──────────────────────────────────────────────────────────────────────
heading2('AI Integration')

heading3('Claude API (Anthropic)')
body('מה זה: ממשק לשימוש במודל השפה Claude.')
body('המודל: claude-sonnet-4-20250514 — מהיר וחכם.')
body('לשם מה: יצירת סיכומים שבועיים + צ\'אט AI מאמן.')

heading3('Edge Function — claude-proxy')
body('למה צריך proxy?')
code_block([
    'בעיה:   Claude API Key חייב להיות סודי.',
    '        אם נשים אותו ב-HTML → כל אחד יכול לגנוב אותו.',
    '',
    'פתרון:  Edge Function (קוד בשרת) שומר את ה-key מוסתר.',
    '        Frontend שולח לEdge Function → היא שולחת לClaude.',
])
body('הזרימה:')
code_block([
    'Browser → Edge Function (Supabase) → Claude API',
    '                    ↑',
    '           רק כאן יש API Key — מוסתר מהמשתמש',
])

heading3('GitHub Pages — Hosting')
body('מה זה: שירות חינמי של GitHub שמגיש קבצי HTML/CSS/JS ישירות מה-Repository.')
body('מגבלה: רק קבצים סטטיים — אין שרת אמיתי. לכן נדרש Supabase לכל לוגיקת ה-Backend.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — Project Structure
# ══════════════════════════════════════════════════════════════════════════════
heading1('3.  מבנה הפרויקט')
divider()

code_block([
    'courtIQ/',
    '│',
    '├── index.html                    ← דף נחיתה שיווקי',
    '│   ├── Navbar + Hero section',
    '│   ├── Stats, How It Works',
    '│   ├── Pricing (3 תוכניות)',
    '│   ├── Features showcase',
    '│   └── Footer',
    '│',
    '├── app.html                      ← האפליקציה הראשית',
    '│   ├── Sidebar (ניווט צדדי)',
    '│   ├── Topbar (מידע עליון)',
    '│   └── 10 פאנלים (ראה סעיף 6)',
    '│',
    '├── sql.html                      ← כלי עזר — review SQL queries',
    '│',
    '├── COURTIQ_GUIDE.md              ← מדריך זה (גרסת Markdown)',
    '│',
    '└── supabase/',
    '    ├── migrations/',
    '    │   └── 001_initial_schema.sql    ← יצירת 3 הטבלאות',
    '    ├── functions/',
    '    │   └── claude-proxy/',
    '    │       └── index.ts              ← Edge Function לClaude',
    '    └── trigger_only.sql              ← Trigger להרשמה אוטומטית',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — Database
# ══════════════════════════════════════════════════════════════════════════════
heading1('4.  מסד הנתונים')
divider()

heading2('טבלה 1: profiles — מי הוא המשתמש')
code_block([
    'profiles',
    '├── id           (UUID)       ← מזהה ייחודי = זהה למזהה ב-auth.users',
    '├── first_name   (text)       ← שם פרטי',
    '├── last_name    (text)       ← שם משפחה',
    '├── position     (text)       ← PG / SG / SF / PF / C',
    '├── skill_level  (text)       ← Beginner / Intermediate / Advanced',
    '├── goal         (text)       ← מטרת האימון',
    '├── plan         (text)       ← starter / pro / elite',
    '├── streak       (int)        ← כמה ימים ברצף אימן',
    '├── created_at   (timestamp)',
    '└── updated_at   (timestamp)',
])

heading2('טבלה 2: training_weeks — שבוע שהסתיים + סיכום AI')
code_block([
    'training_weeks',
    '├── id           (UUID)       ← מזהה ייחודי לשבוע',
    '├── user_id      (UUID)       ← Foreign Key ← auth.users',
    '├── week_number  (int)        ← מספר השבוע (1, 2, 3…)',
    '├── label        (text)       ← תווית מוצגת — "W4"',
    '├── summary_json (jsonb)      ← כל ניתוח ה-AI שמור כ-JSON',
    '└── created_at   (timestamp)',
])

body('מה summary_json מכיל:')
code_block([
    '{',
    '  "averages":   { "shootingPct": 67.3, "verticalIn": 28.5, "sprintSec": 4.2 },',
    '  "strengths":  ["Shooting consistency", "Sprint speed"],',
    '  "focusAreas": ["Vertical jump needs work"],',
    '  "feedback":   "טקסט AI מפורט על השבוע...",',
    '  "generatedAt": "2025-03-01T10:00:00Z"',
    '}',
])

heading2('טבלה 3: training_sessions — אימון יומי בודד')
code_block([
    'training_sessions',
    '├── id               (UUID)',
    '├── user_id          (UUID)    ← Foreign Key ← auth.users',
    '├── week_id          (UUID)    ← Foreign Key ← training_weeks (אופציונלי)',
    '├── day              (text)    ← "Mon", "Tue", "Wed"…',
    '├── shots_made       (numeric) ← זריקות שנכנסו',
    '├── shots_attempted  (numeric) ← זריקות שנזרקו',
    '├── dribbling_min    (numeric) ← דקות דריבלינג',
    '├── vertical_in      (numeric) ← קפיצה אנכית באינצ\'ים',
    '├── sprint_sec       (numeric) ← זמן ספרינט בשניות',
    '├── intensity        (text)    ← Low / Medium / High',
    '├── notes            (text)    ← הערות חופשיות',
    '└── created_at       (timestamp)',
])

heading2('Foreign Keys — מה זה ולמה חשוב')
body('דמיה: כמו בטופס בנק — עמודת "קוד סניף" מצביעה על טבלת "סניפים".')
code_block([
    'training_sessions.user_id  →  auth.users.id',
    '      ↑',
    'אם מוחקים משתמש → כל הסשנים שלו נמחקים אוטומטית (CASCADE)',
])

heading2('Trigger — יצירת פרופיל אוטומטית')
body('מה זה Trigger: קוד SQL שרץ אוטומטית כתגובה לאירוע.')
body('למה צריך: הרשמה יוצרת שורה ב-auth.users (פנימי ל-Supabase) אבל לא ב-profiles שלנו. ה-Trigger גשר על הפער.')
code_block([
    '-- כשמשתמש חדש נרשם:',
    'create trigger on_auth_user_created',
    '  after insert on auth.users',
    '  for each row',
    '  execute procedure handle_new_user();',
    '',
    '-- handle_new_user() יוצרת שורה ב-profiles עם id + שם.',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — AI Integration
# ══════════════════════════════════════════════════════════════════════════════
heading1('5.  האינטגרציה עם Claude AI')
divider()

heading2('זרימת יצירת הסיכום השבועי')

heading3('שלב 1 — איסוף נתונים מה-DB')
code_block([
    'const { data: sessions } = await supabase',
    '  .from("training_sessions")',
    '  .select("*")',
    '  .eq("user_id", APP.user.id)',
    '  .gte("created_at", weekStart)',
])

heading3('שלב 2 — בניית Prompt')
code_block([
    'const systemPrompt = `You are an elite basketball performance analyst.',
    r'The athlete plays ${APP.profile.position}.`',
    '',
    'const userMessage = `Sessions this week:',
    r'  Monday: ${mon.shots_made}/${mon.shots_attempted} shots,',
    r'           vertical ${mon.vertical_in}"`',
])

heading3('שלב 3 — קריאה לEdge Function')
code_block([
    'const response = await fetch(',
    '  "https://[project].supabase.co/functions/v1/claude-proxy",',
    '  {',
    '    method: "POST",',
    '    headers: {',
    '      Authorization: `Bearer ${session.access_token}`,',
    '      "Content-Type": "application/json"',
    '    },',
    '    body: JSON.stringify({ messages, system: systemPrompt })',
    '  }',
    ')',
])

heading3('שלב 4 — שמירה בDB')
code_block([
    'await supabase.from("training_weeks").insert({',
    '  user_id:     APP.user.id,',
    '  week_number: currentWeek,',
    '  summary_json: aiResult',
    '})',
])

heading2('System vs User Messages — ההבדל')
make_table(
    ['סוג', 'מה זה', 'דוגמה'],
    [
        ['system', 'הוראות קבועות לAI — מי הוא, מה תפקידו', '"You are an elite basketball coach"'],
        ['user',   'הבקשה הספציפית — הנתונים לניתוח',         '"Here are this week\'s sessions: ..."'],
    ],
    col_widths=[1.2, 2.8, 2.2]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — Features
# ══════════════════════════════════════════════════════════════════════════════
heading1('6.  כל תכונה שנבנתה — הסבר מפורט')
divider()

heading2('תוכניות המנוי')
make_table(
    ['תכונה', 'Starter $9', 'Pro $24', 'Elite $59'],
    [
        ['AI תוכנית שבועית',      '✅', '✅', '✅'],
        ['ספריית תרגילים',         '50+', '500+', '500+'],
        ['מעקב ביצועים',           'בסיסי', 'מלא', 'מלא'],
        ['Video Drills',           '❌', '✅', '✅'],
        ['AI Coach Chat',          '❌', '50 הודעות', 'ללא הגבלה'],
        ['Game Film Analysis',     '❌', '❌', '✅'],
        ['Nutrition & Recovery',   '❌', '❌', '✅'],
    ],
    col_widths=[2.5, 1.2, 1.2, 1.2]
)

panels = [
    ("פאנל 1: Today's Program",
     "תוכנית האימון היומית, לפי הפוזיציה.",
     ["תרגילים מסוננים לפי position", "KPIs: streak ימים, שבוע נוכחי, % זריקה, קפיצה"]),

    ("פאנל 2: Log Session",
     "טופס לרישום אימון שהסתיים.",
     ["יום האימון", "זריקות שנכנסו / זריקות שנזרקו (מחשב % אוטומטי)",
      "דקות דריבלינג, קפיצה אנכית, זמן ספרינט, עוצמה, הערות",
      "אחרי שמירה: INSERT לטבלת training_sessions"]),

    ("פאנל 3: Drill Library",
     "ספריית 50+ תרגילים עם סינון.",
     ["פילטרים: קטגוריה (Shooting/Dribbling/Defense) + קושי",
      "מגבלה: נתוני תרגילים hardcoded בJS — לא מDB"]),

    ("פאנל 4: Video Drills (PRO)",
     "ספריית וידאו HD של תרגילים.",
     ["ממשק קיים, אין תוכן אמיתי — Feature gated"]),

    ("פאנל 5: Analytics",
     "גרפים וניתוח ביצועים.",
     ["4 גרפים Chart.js — shooting%, vertical, sprint, השוואה שבועית",
      "AI Summary — סיכום שבועי שנוצר על ידי Claude",
      "Strengths & Focus Areas"]),

    ("פאנל 6: History",
     "ארכיון שבועות אימון.",
     ["רשימת שבועות עם תאריך ו-KPIs בסיסיים"]),

    ("פאנל 7: Game Film (ELITE)",
     "העלאת סרטוני משחק לניתוח AI.",
     ["UI קיים, אין backend — Feature gated"]),

    ("פאנל 8: AI Coach Chat (PRO)",
     "צ'אט בזמן אמת עם מאמן AI.",
     ["שאלה → Edge Function → Claude (עם context מלא) → תשובה",
      "מגבלה: היסטוריה נמחקת ברענון — לא שמורה בDB"]),

    ("פאנל 9: Weekly Calendar",
     "לוח שנה שבועי לתכנון אימונים.",
     ["UI בסיסי, אין לוגיקה"]),

    ("פאנל 10: Settings",
     "ניהול פרופיל.",
     ["שינוי שם, פוזיציה, רמה, מטרה, תוכנית",
      "אחרי שמירה: UPDATE לטבלת profiles"]),
]

for title, desc, bullets_list in panels:
    heading3(title)
    body(desc)
    for b in bullets_list:
        bullet(b)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — What's Missing
# ══════════════════════════════════════════════════════════════════════════════
heading1('7.  מה חסר — ניתוח מלא')
divider()

heading2('בעיות אבטחה')
make_table(
    ['בעיה', 'חומרה', 'הסבר'],
    [
        ['אין auth guard בapp.html', '🔴 קריטי',  'כל אחד יכול להיכנס ללא התחברות'],
        ['innerHTML ללא sanitization', '🟠 גבוהה', 'סיכון XSS — הזרקת קוד זדוני'],
        ['API keys נגלים בHTML', '🟡 בינונית', 'Supabase anon key נחשף (מקובל אך שים לב)'],
    ],
    col_widths=[2.5, 1.4, 2.3]
)

heading3('XSS מוסבר')
code_block([
    'תוקף שולח הודעת צ\'אט: "<script>sendAllData(hackerSite)</script>"',
    'אם משתמשים ב-innerHTML ישירות → הסקריפט מופעל',
    'פתרון: textContent במקום innerHTML, או ספריית DOMPurify',
])

heading2('תכונות שמופיעות אבל לא פועלות')
make_table(
    ['תכונה', 'מה קיים', 'מה חסר'],
    [
        ['Video Drills',         'ממשק יפה',   'וידאו אמיתי, storage בSupabase'],
        ['Game Film Analysis',   'Upload UI',   'backend, AI analysis pipeline'],
        ['Nutrition & Recovery', 'Mock data',   'AI generation, שמירה בDB'],
        ['Calendar',             'UI בסיסי',    'לוגיקת לוח שנה, scheduling'],
        ['Notifications',        'פאנל ריק',    'כל המערכת'],
    ],
    col_widths=[2.0, 1.8, 2.4]
)

heading2('בעיות טכניות')
make_table(
    ['בעיה', 'השפעה'],
    [
        ['אין error handling מסודר', 'המשתמש לא יודע כשמשהו נכשל'],
        ['Chat history ב-memory בלבד', 'נאבד ברענון הדף'],
        ['אין validation בטפסים', 'ניתן לשלוח נתונים ריקים'],
        ['הכל בקובץ HTML אחד', 'קשה לתחזוקה ועדכונים עתידיים'],
    ],
    col_widths=[2.8, 3.4]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — Roadmap
# ══════════════════════════════════════════════════════════════════════════════
heading1('8.  מפת דרכים לאפליקציה המושלמת')
divider()

# Phase 1
heading2('Phase 1 — יסודות  (2–3 שבועות)')
body('מטרה: לוודא שמה שקיים עובד כמו שצריך.')

heading3('1.1  Auth Guard — הגנה על הדשבורד')
code_block([
    '// להוסיף לתחילת app.html:',
    'const { data: { session } } = await supabase.auth.getSession()',
    'if (!session) {',
    '  window.location.href = "/index.html#login"',
    '}',
])

heading3('1.2  Error Handling מסודר')
code_block([
    'function showError(message) {',
    '  const toast = document.createElement("div")',
    '  toast.className = "toast toast-error"',
    '  toast.textContent = message',
    '  document.body.appendChild(toast)',
    '  setTimeout(() => toast.remove(), 4000)',
    '}',
    '',
    '// שימוש:',
    'try {',
    '  await saveSession(data)',
    '} catch (err) {',
    '  showError("שגיאה בשמירת האימון. נסה שוב.")',
    '}',
])

heading3('1.3  Chat History Persistence')
code_block([
    'create table chat_messages (',
    '  id         uuid primary key default gen_random_uuid(),',
    '  user_id    uuid references auth.users on delete cascade,',
    '  role       text not null,    -- "user" / "assistant"',
    '  content    text not null,',
    '  created_at timestamptz default now()',
    ');',
])

heading3('1.4  Input Sanitization — XSS Protection')
code_block([
    '// במקום innerHTML:',
    'element.innerHTML = userContent          // ← סכנה!',
    '',
    '// תשתמש ב-textContent:',
    'element.textContent = userContent        // ← בטוח',
    '',
    '// או ספריית DOMPurify:',
    'element.innerHTML = DOMPurify.sanitize(userContent)',
])

# Phase 2
heading2('Phase 2 — תכונות ליבה  (3–5 שבועות)')
body('מטרה: להגיש את מה שמובטח בPRO ו-ELITE.')

heading3('2.1  Video Drill Library (PRO)')
body('מה צריך:')
bullet('Storage: Supabase Storage bucket לאחסון וידאו')
bullet('Database: טבלת drills עם שדה video_url')
bullet('Frontend: Video player component (HTML5 <video> tag)')
code_block([
    'create table drills (',
    '  id            uuid primary key default gen_random_uuid(),',
    '  title         text not null,',
    '  category      text not null,',
    '  difficulty    text not null,',
    '  video_url     text,        ← URL לSupabase Storage',
    '  plan_required text default "starter"',
    ');',
])

heading3('2.2  Game Film Analysis (ELITE)')
body('זרימה:')
code_block([
    'משתמש מעלה וידאו → Supabase Storage',
    '→ Edge Function → שולח לClaude עם הוראות ניתוח',
    '→ Claude מחזיר analysis → נשמר בDB',
    '→ מוצג למשתמש כדוח (ניתן לייצוא PDF)',
])

heading3('2.3  Nutrition & Recovery (ELITE)')
bullet('שאלון התחלתי: משקל, גובה, יעד קלורי')
bullet('Edge Function שמייצרת תפריט שבועי עם Claude')
bullet('שמירה בDB לכל שבוע')

heading3('2.4  Calendar עם AI Scheduling')
bullet('ספריית לוח שנה (Flatpickr)')
bullet('לוגיקת scheduling — מנוחה בין אימונים, אי-חפיפה')
bullet('AI מציע מתי לאמן לפי עומס השבוע')

# Phase 3
heading2('Phase 3 — פולישינג  (2–3 שבועות)')
body('מטרה: חווית משתמש מלאה ומלוטשת.')

heading3('3.1  Push Notifications')
code_block([
    'const registration = await navigator.serviceWorker.register("/sw.js")',
    'const subscription = await registration.pushManager.subscribe({',
    '  userVisibleOnly: true,',
    '  applicationServerKey: VAPID_KEY',
    '})',
    '// שמירת subscription בSupabase לשליחת התראות',
])

heading3('3.2  Data Export (CSV / PDF)')
code_block([
    'function exportToCSV(sessions) {',
    '  const csv = sessions.map(s =>',
    '    `${s.day},${s.shots_made},${s.shots_attempted},${s.vertical_in}`',
    '  ).join("\\n")',
    '  const blob = new Blob([csv], { type: "text/csv" })',
    '  // trigger download...',
    '}',
])

heading3('3.3  Progressive Web App (PWA)')
body('מה זה: הופך את האתר לאפליקציה שניתן "להתקין" בטלפון.')
code_block([
    '// manifest.json',
    '{',
    '  "name":             "CourtIQ",',
    '  "short_name":       "CourtIQ",',
    '  "theme_color":      "#f5a623",',
    '  "background_color": "#0c0d0f",',
    '  "display":          "standalone",',
    '  "start_url":        "/app.html"',
    '}',
])

# Phase 4
heading2('Phase 4 — Scale  (חודש+)')
body('מטרה: להפוך לעסק אמיתי.')

heading3('4.1  Payment Integration — Stripe')
code_block([
    'const stripe  = Stripe(STRIPE_PUBLIC_KEY)',
    'const session = await createCheckoutSession({',
    '  priceId: "price_pro_monthly",',
    '  userId:  APP.user.id',
    '})',
    'window.location.href = session.url',
    '',
    '// Webhook: Stripe → Supabase Edge Function → עדכון plan בDB',
])

heading3('4.2  Admin Panel')
bullet('ניהול משתמשים')
bullet('ניהול תוכן — תרגילים, וידאו')
bullet('Analytics — מספר משתמשים, revenue, churn')

heading3('4.3  Social Features')
bullet('Leaderboard — דירוג שחקנים')
bullet('Training Challenges — תחרויות שבועיות')
bullet('Coach Assignment — שיוך מאמן אנושי')

heading3('4.4  Native Mobile App')
make_table(
    ['אפשרות', 'גישה', 'יתרון'],
    [
        ['React Native', 'קוד אחד לiOS ו-Android', 'ביצועים גבוהים'],
        ['Capacitor',    'עוטף את האתר הנוכחי',    'הכי מהיר — קוד כבר קיים'],
    ],
    col_widths=[1.5, 2.5, 2.2]
)
body('המלצה: Capacitor — כי הקוד כבר קיים ואין צורך לכתוב מחדש.', indent=True)

# Phase 5
heading2('Phase 5 — AI Upgrade')
body('מטרה: AI ברמה הבאה.')

heading3('5.1  Pose Estimation — ניתוח תנועה')
code_block([
    'מצלמה → MediaPipe (Google) → ניתוח מפרקים',
    '→ feedback אוטומטי על פורמת הזריקה',
    '',
    'ספרייה: Google MediaPipe — חינמי, עובד ישירות בBrowser',
])

heading3('5.2  AI Personalized — מתחכם עם הזמן')
code_block([
    '// במקום prompt גנרי:',
    '"נתח את הנתונים האלה..."',
    '',
    '// Prompt אישי שמשתמש בהיסטוריה:',
    '"שחקן זה נוטה להתמוטט בדיוק זריקה תחת לחץ.',
    ' בשבועות ראשונים: 65%, ירד ל-55% תחת לחץ.',
    ' תן המלצות ספציפיות לתסריט זה..."',
])

heading3('5.3  Streaming Responses — תגובה בזמן אמת')
code_block([
    '// במקום לחכות לתשובה מלאה — טוקן אחר טוקן:',
    'const stream = await fetch(edgeFunctionUrl, { ... })',
    'const reader = stream.body.getReader()',
    '',
    'while (true) {',
    '  const { done, value } = await reader.read()',
    '  if (done) break',
    '  chatBox.textContent += new TextDecoder().decode(value)',
    '}',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — Glossary
# ══════════════════════════════════════════════════════════════════════════════
heading1('9.  מילון מונחים')
divider()

glossary = [
    ('API',              'Application Programming Interface — ממשק לתקשורת בין מערכות'),
    ('Auth',             'Authentication — תהליך זיהוי משתמש (כניסה)'),
    ('Backend',          'צד השרת — קוד שרץ בענן, לא בדפדפן'),
    ('CDN',              'Content Delivery Network — רשת שרתים לטעינה מהירה'),
    ('CORS',             'Cross-Origin Resource Sharing — מדיניות אבטחה בין דומיינים'),
    ('CRUD',             'Create, Read, Update, Delete — 4 פעולות בסיסיות על DB'),
    ('CSS Variables',    'משתנים ב-CSS שאפשר לשתף בין חלקים'),
    ('Deno',             'Runtime לJavaScript בצד שרת (כמו Node.js, אבל מאובטח יותר)'),
    ('DOM',              'Document Object Model — ייצוג HTML כעצים, ניתן לשינוי בJS'),
    ('Edge Function',    'קוד שרץ "קרוב" למשתמש, על שרתים מבוזרים בעולם'),
    ('Foreign Key',      'עמודה בטבלה שמפנה לשורה בטבלה אחרת'),
    ('Frontend',         'צד הלקוח — HTML/CSS/JS שרץ בדפדפן'),
    ('Git',              'מערכת ניהול גרסאות — tracking של שינויים בקוד'),
    ('GitHub Pages',     'שירות חינמי להגשת אתרים סטטיים'),
    ('Hardcoded',        'ערכים קבועים ישירות בקוד, לא מDB'),
    ('innerHTML',        'דרך לשנות תוכן HTML דרך JS — עלולה לגרום XSS'),
    ('JWT',              'JSON Web Token — תעודת זהות דיגיטלית מוצפנת'),
    ('JSONB',            'סוג עמודה בPostgreSQL לאחסון JSON הניתן לחיפוש'),
    ('Migration',        'קובץ SQL שמגדיר שינויים למבנה ה-DB'),
    ('Mock Data',        'נתונים מדומים לצורכי פיתוח ו-UI'),
    ('Node.js',          'Runtime לJavaScript בצד שרת'),
    ('Payload',          'הנתונים שנשלחים ב-API request'),
    ('PostgreSQL',       'מסד נתונים רלציוני מתקדם, בסיס Supabase'),
    ('Prompt',           'ההוראות שנשלחות למודל AI'),
    ('PWA',              'Progressive Web App — אתר שמתנהג כApp נייד'),
    ('RLS',              'Row Level Security — אבטחה ברמת שורה ב-DB'),
    ('Repository',       'תיקיה עם כל הקוד, מנוהלת על ידי Git'),
    ('Runtime',          'הסביבה שמריצה קוד (דפדפן, Node, Deno)'),
    ('Session',          'מצב כניסה פעיל של משתמש'),
    ('State',            'המצב הנוכחי של האפליקציה בזיכרון'),
    ('Static Site',      'אתר מורכב מקבצים סטטיים בלבד, ללא שרת'),
    ('Supabase',         'Backend-as-a-Service — מסד נתונים + Auth + Edge Functions'),
    ('Token',            'מחרוזת מוצפנת שמשמשת לאימות'),
    ('Trigger',          'קוד SQL שמופעל אוטומטית כתגובה לאירוע ב-DB'),
    ('TypeScript',       'JavaScript עם type checking — מונע שגיאות'),
    ('UUID',             'Universally Unique Identifier — מזהה ייחודי אוניברסלי'),
    ('Vanilla JS',       'JavaScript ללא frameworks כמו React / Vue'),
    ('Webhook',          'HTTP request שנשלח אוטומטית כשאירוע מתרחש (למשל תשלום)'),
    ('XSS',              'Cross-Site Scripting — התקפת הזרקת קוד זדוני'),
]

make_table(
    ['מונח', 'הסבר'],
    glossary,
    col_widths=[1.8, 4.4]
)

# ─── Final Summary ────────────────────────────────────────────────────────────
doc.add_page_break()
heading1('סיכום — מפת הדרך המצומצמת')
divider()

code_block([
    'עכשיו:          auth guard + error handling + chat persistence',
    '1-2 חודשים:    video drills + nutrition + calendar',
    '3-4 חודשים:    payment (Stripe) + PWA + notifications',
    '5-6 חודשים:    mobile app + social features + admin panel',
    'שנה+:           pose estimation + AI personalization + coaching platform',
])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CourtIQ  ·  חוברת מקצועית  ·  מרץ 2026')
run.font.size      = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic         = True

# ─── Save ────────────────────────────────────────────────────────────────────
doc.save('/home/user/courtIQ/CourtIQ_Guide.docx')
print('Done → CourtIQ_Guide.docx')
